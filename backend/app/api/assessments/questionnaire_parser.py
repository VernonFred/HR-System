"""问卷导入解析模块.

支持从多种格式导入问卷：
- JSON 文件
- Excel 文件 (.xlsx)
- Word 文件 (.docx)
- 纯文本文件 (.txt)

V45: 新增AI智能解析功能
- 优先使用AI识别题目类型和选项
- 规则匹配作为兜底方案
"""
import json
import re
import logging
import asyncio
from typing import Dict, List, Any, Optional, Tuple

logger = logging.getLogger(__name__)

from app.api.assessments.questionnaire_text_parser import (
    _add_option_to_question,
    _create_question_from_text,
    _looks_like_question,
    _parse_text,
    _post_process_questions,
)

# ========== V45: AI智能解析 ==========

# AI解析提示词（V45优化版 - 基于ChatGPT建议）
AI_PARSE_PROMPT = """你是一名"问卷解析助手"，需要把一整份问卷的原始文本解析成结构化 JSON。

【输入说明】
- 输入是一段从 Word / 网页 / PDF 中复制出来的问卷文本。
- 里面可能包含：问卷标题、说明文字、分节标题（如"一、基本信息"）、题号（1. / 2. / （1））、题型标注（单选题、多选题、量表题、是非题、填空题等）。
- 每道题下面可能有 A/B/C/D 等选项，也可能是开放题没有选项。

【你的任务】
1. 找出问卷标题（如果有），放在 "title" 字段里；如果没有标题，title 用空字符串 ""。
2. 只提取真正的题目，忽略以下内容：
   - 问卷开头或结尾的说明文字（如"感谢填写本问卷"）。
   - 分节标题（如"第一部分 基本信息"、"一、学习情况"）。
   - 页码、装饰性文字等。
3. 对每一道题，输出：
   - text：题目正文，不要带题号，不要带"（单选题）"等标签。
   - type：题目类型，只能是以下几种之一：
       - "single"   单选题（只能选一个）
       - "multiple" 多选题（可以选多个）
       - "rating"   量表/打分题（如 1–5 分、"非常不同意–非常同意"）
       - "yesno"    是非题 / 判断题（例如"是/否"、"对/错"）
       - "choice"   二选一题（两个较长的选项，如A/B两种观点）
       - "text"     文本开放题（填空、简答，没有固定选项）
       - "textarea" 多行文本题（需要详细描述的开放题）
   - options：一个对象数组，每个对象包含 text 和 score 字段：
       - 对于 single / multiple / yesno / choice，填入所有选项；
       - 对于 rating，如果题目给出了文字锚点（如"1 非常不同意…5 非常同意"），也列出；
       - 对于 text / textarea，没有选项时，options 用空数组 []。
   - scale（可选）：如果是 rating 题目，增加量表范围信息：
       - "scale": {{"min": 1, "max": 5, "minLabel": "非常不同意", "maxLabel": "非常同意"}}

【题型识别规则（请严格遵守）】
1. 如果题干或括号里明确写了"单选题""单选""请选择一项"，则 type = "single"。
2. 如果题干或括号里写了"多选题""多选""可多选""至少选择两项"等，则 type = "multiple"。
3. 如果出现"在 1–5 分中选择""请按 1~7 分打分""非常不同意 – 非常同意"等评分/量表描述，则 type = "rating"。
4. 如果选项只有"是/否""对/错""是的/不是"，或者题干里写了"是否……"，并且没有更复杂的选项，则 type = "yesno"。
5. 如果只有两个选项，且每个选项内容较长（超过15个字），则 type = "choice"。
6. 如果是"简要说明……""请填写……""其他情况请写出""请描述……"，且没有选项：
   - 如果需要详细描述（如"请详细说明"），type = "textarea"；
   - 其他简短填写，type = "text"。
7. 选项前面的序号或字母（如"1."、"A."、"B、"等）请去掉，只保留选项内容本身。

【输出格式要求】
- 严格输出一个 JSON 对象，不要包含任何额外说明文字，不要使用 Markdown 代码块。
- 字段结构固定为：

{{
  "title": "问卷标题（字符串）",
  "description": "问卷描述（如果有）",
  "questions": [
    {{
      "text": "题目内容（字符串）",
      "type": "single/multiple/rating/yesno/choice/text/textarea 之一",
      "options": [
        {{"text": "选项内容", "score": 0}}
      ],
      "scale": {{
        "min": 1,
        "max": 5,
        "minLabel": "最低标签",
        "maxLabel": "最高标签"
      }}
    }}
  ]
}}

【待解析的问卷原文】
{content}"""


async def _ai_parse_content(content: str) -> Optional[Dict[str, Any]]:
    """使用AI解析问卷内容."""
    try:
        # 导入AI客户端
        from app.core.ai.portrait_router import call_portrait_model
        
        prompt = AI_PARSE_PROMPT.format(content=content[:8000])  # 限制内容长度
        
        messages = [
            {"role": "system", "content": "你是一个专业的问卷解析助手，只输出JSON格式的结果。"},
            {"role": "user", "content": prompt}
        ]
        
        logger.info("🤖 开始AI智能解析问卷...")
        
        result = await call_portrait_model(
            messages=messages,
            level="normal",  # 使用普通级别即可
            max_tokens=4096,
            temperature=0.1,  # 低温度确保稳定输出
        )
        
        # 提取AI返回的内容
        ai_content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        
        if not ai_content:
            logger.warning("AI返回内容为空")
            return None
        
        # 尝试提取JSON
        json_match = re.search(r'```json\s*(.*?)\s*```', ai_content, re.DOTALL)
        if json_match:
            json_str = json_match.group(1)
        else:
            # 尝试直接解析
            json_str = ai_content.strip()
            # 移除可能的前后缀
            if json_str.startswith("```"):
                json_str = re.sub(r'^```\w*\n?', '', json_str)
                json_str = re.sub(r'\n?```$', '', json_str)
        
        parsed = json.loads(json_str)
        logger.info(f"✅ AI解析成功，识别到 {len(parsed.get('questions', []))} 道题目")
        return parsed
        
    except json.JSONDecodeError as e:
        logger.warning(f"AI返回的JSON解析失败: {e}")
        return None
    except Exception as e:
        logger.warning(f"AI解析失败: {e}")
        return None


def _convert_ai_result_to_questions(ai_result: Dict[str, Any]) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """将AI解析结果转换为标准格式."""
    metadata = {
        "name": ai_result.get("title") or "导入的问卷",
        "description": ai_result.get("description") or "",
        "estimated_minutes": 15,
    }
    
    questions = []
    for i, q in enumerate(ai_result.get("questions", [])):
        q_type = q.get("type", "single")
        
        question = {
            "id": f"q{i+1}",
            "text": q.get("text", ""),
            "type": q_type,
            "options": [],
            "required": True,
        }
        
        # 处理选项
        for j, opt in enumerate(q.get("options", [])):
            if isinstance(opt, str):
                question["options"].append({
                    "id": f"q{i+1}_opt{j}",
                    "text": opt,
                    "score": 0
                })
            elif isinstance(opt, dict):
                question["options"].append({
                    "id": f"q{i+1}_opt{j}",
                    "text": opt.get("text", ""),
                    "score": opt.get("score", 0)
                })
        
        # V45: 处理量表题的scale信息
        if q_type == "rating" and q.get("scale"):
            scale = q.get("scale")
            question["scale"] = {
                "min": scale.get("min", 1),
                "max": scale.get("max", 5),
                "minLabel": scale.get("minLabel", "最低"),
                "maxLabel": scale.get("maxLabel", "最高"),
            }
        
        questions.append(question)
    
    return metadata, questions


async def parse_questionnaire_file_async(
    content: bytes,
    filename: str,
    content_type: str,
    use_ai: bool = True
) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """
    异步解析问卷文件（支持AI智能解析）.
    
    Args:
        content: 文件内容
        filename: 文件名
        content_type: 内容类型
        use_ai: 是否使用AI解析（默认True）
    
    Returns:
        Tuple[问卷元数据, 题目列表]
    """
    # 先提取文本内容
    text_content = _extract_text_content(content, filename)
    
    # V45: 优先尝试AI解析
    if use_ai and text_content:
        try:
            ai_result = await _ai_parse_content(text_content)
            if ai_result and ai_result.get("questions"):
                logger.info("✅ 使用AI解析结果")
                return _convert_ai_result_to_questions(ai_result)
        except Exception as e:
            logger.warning(f"AI解析异常，使用规则匹配: {e}")
    
    # 兜底：使用规则匹配
    logger.info("📋 使用规则匹配解析")
    return parse_questionnaire_file(content, filename, content_type)


def _extract_text_content(content: bytes, filename: str) -> str:
    """从文件中提取纯文本内容（供AI解析使用）."""
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith('.json'):
            return content.decode('utf-8')
        
        elif filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls'):
            try:
                import openpyxl
                from io import BytesIO
                wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
                ws = wb.active
                lines = []
                for row in ws.iter_rows(values_only=True):
                    if row and row[0]:
                        lines.append(str(row[0]).strip())
                return '\n'.join(lines)
            except Exception:
                return ""
        
        elif filename_lower.endswith('.docx'):
            try:
                from docx import Document
                from io import BytesIO
                doc = Document(BytesIO(content))
                return '\n'.join(para.text.strip() for para in doc.paragraphs if para.text.strip())
            except Exception:
                return ""
        
        elif filename_lower.endswith('.txt'):
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('gbk', errors='ignore')
        
        return ""
    except Exception as e:
        logger.warning(f"提取文本内容失败: {e}")
        return ""


def parse_questionnaire_file(
    content: bytes,
    filename: str,
    content_type: str
) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """
    解析上传的问卷文件.
    
    Returns:
        Tuple[问卷元数据, 题目列表]
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.json'):
        return _parse_json(content)
    elif filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls'):
        return _parse_excel(content)
    elif filename_lower.endswith('.docx'):
        return _parse_word(content)
    elif filename_lower.endswith('.txt'):
        return _parse_text(content)
    else:
        raise ValueError(f"不支持的文件格式: {filename}")


def _parse_json(content: bytes) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """解析JSON格式问卷."""
    try:
        data = json.loads(content.decode('utf-8'))
    except json.JSONDecodeError as e:
        raise ValueError(f"JSON格式错误: {str(e)}")
    
    # 支持多种JSON结构
    if isinstance(data, list):
        # 直接是题目列表
        return {}, _normalize_questions(data)
    
    if isinstance(data, dict):
        # 标准结构: { "title": "...", "questions": [...] }
        metadata = {
            "name": data.get("title") or data.get("name") or "导入的问卷",
            "description": data.get("description") or data.get("desc") or "",
            "estimated_minutes": data.get("estimated_minutes") or data.get("duration") or 15,
        }
        
        questions = data.get("questions") or data.get("items") or []
        return metadata, _normalize_questions(questions)
    
    raise ValueError("无法识别的JSON结构")


def _parse_excel(content: bytes) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """解析Excel格式问卷."""
    try:
        import openpyxl
        from io import BytesIO
    except ImportError:
        raise ValueError("需要安装openpyxl库来解析Excel文件")
    
    wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
    ws = wb.active
    
    questions = []
    metadata = {"name": "导入的问卷", "description": "", "estimated_minutes": 15}
    
    # 尝试从第一行获取标题
    first_row = [cell.value for cell in ws[1]]
    if first_row and first_row[0] and not _looks_like_question(str(first_row[0])):
        metadata["name"] = str(first_row[0])
        start_row = 2
    else:
        start_row = 1
    
    # 解析题目
    current_question = None
    for row in ws.iter_rows(min_row=start_row, values_only=True):
        if not row or not row[0]:
            continue
        
        text = str(row[0]).strip()
        if not text:
            continue
        
        # 判断是题目还是选项
        if _looks_like_question(text):
            if current_question:
                questions.append(current_question)
            current_question = _create_question_from_text(text, len(questions) + 1)
        elif current_question and _looks_like_option(text):
            _add_option_to_question(current_question, text)
    
    if current_question:
        questions.append(current_question)
    
    # V45: 后处理 - 根据选项智能推断题目类型
    questions = _post_process_questions(questions)
    
    return metadata, questions


def _parse_word(content: bytes) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """解析Word格式问卷."""
    try:
        from docx import Document
        from io import BytesIO
    except ImportError:
        raise ValueError("需要安装python-docx库来解析Word文件")
    
    doc = Document(BytesIO(content))
    
    metadata = {"name": "导入的问卷", "description": "", "estimated_minutes": 15}
    questions = []
    current_question = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 第一段可能是标题
        if not questions and not current_question and not _looks_like_question(text):
            metadata["name"] = text
            continue
        
        if _looks_like_question(text):
            if current_question:
                questions.append(current_question)
            current_question = _create_question_from_text(text, len(questions) + 1)
        elif current_question and _looks_like_option(text):
            _add_option_to_question(current_question, text)
    
    if current_question:
        questions.append(current_question)
    
    # V45: 后处理 - 根据选项智能推断题目类型
    questions = _post_process_questions(questions)
    
    return metadata, questions


def _normalize_questions(questions: List[Any]) -> List[Dict[str, Any]]:
    """标准化题目列表格式."""
    normalized = []
    
    for i, q in enumerate(questions):
        if isinstance(q, str):
            # 纯文本题目
            normalized.append({
                "id": f"q{i+1}",
                "text": q,
                "type": "single",
                "options": [],
                "required": True
            })
        elif isinstance(q, dict):
            # 字典格式题目
            normalized.append({
                "id": q.get("id") or f"q{i+1}",
                "text": q.get("text") or q.get("question") or q.get("title") or "",
                "type": q.get("type") or "single",
                "options": _normalize_options(q.get("options") or q.get("choices") or []),
                "required": q.get("required", True),
                "score": q.get("score", 0)
            })
    
    return normalized


def _normalize_options(options: List[Any]) -> List[Dict[str, Any]]:
    """标准化选项列表格式."""
    normalized = []
    
    for i, opt in enumerate(options):
        if isinstance(opt, str):
            normalized.append({
                "id": f"opt{i}",
                "text": opt,
                "score": 0
            })
        elif isinstance(opt, dict):
            normalized.append({
                "id": opt.get("id") or f"opt{i}",
                "text": opt.get("text") or opt.get("label") or "",
                "score": opt.get("score") or opt.get("value") or 0
            })
    
    return normalized

